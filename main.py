# Receipt maker
# by Justin Rainier Go as part of a Python demonstration

import tkinter as tk
from tkinter import ttk
from tkinter import filedialog
from tkinter import messagebox
from datetime import datetime
from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class App(tk.Frame):
    HELP_MSG = """1. Load item list: select an Excel file to load items from. (See the sample items.xlsx for the template)\n
2. The items will show up on the Item list. Clicking on one of these items will select it; clicking on it again will add 1 item to the receipt.\n
3. If you want to remove an item from the receipt, select it on the Receipt preview and click it again.\n
4. Set the customer name by typing in the Customer name entry box.\n
5. Export receipt: this will export your receipt to a Word file which you can then print."""

    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.pack()
        self.item_dict = {} # {item: price}, intialized by load_items_from_workbook()
        self.receipt_items = {} # {item: amount}, affected by add_item() and remove_item()
        self.receipt_items_for_docx = [] # [(amount, item, price)]
        self.total_price = 0
        self.customer_name = tk.StringVar()
        # Top menu
        self.menu = tk.Menu(self.master)
        self.master.config(menu=self.menu)
        self.menu.add_command(label="Load item list", command=self.load_items_from_workbook)
        self.menu.add_command(label="Export receipt", command=self.export_to_document)
        self.menu.add_command(label="Help", command=self.help)
        
        # Item list
        self.menu_frame = tk.Frame(self)
        self.menu_frame.pack(side=tk.LEFT)
        self.menu_label = tk.Label(self.menu_frame, text="Item list")
        self.menu_label.grid(column=0, row=0)
        self.item_list_display = ttk.Treeview(
            self.menu_frame,
            columns=(1,2),
            show="headings",
            height=10
        )
        self.item_list_display.grid(column=0, row=1)
        self.item_list_display.heading(1, text="Item")
        self.item_list_display.heading(2, text="Price")
        self.item_list_display.bind('<Double-1>', self.add_item)

        # Item search
        self.search_frame = tk.Frame(self.menu_frame)
        self.search_frame.grid(column=0, row=2)
        self.search_label = tk.Label(self.search_frame, text="Search (Not working yet)")
        self.search_label.pack(side=tk.LEFT)
        self.search_entry = tk.Entry(self.search_frame)
        self.search_entry.pack(side=tk.RIGHT)

        # Receipt preview
        self.receipt_frame = tk.Frame(self)
        self.receipt_label = tk.Label(self.receipt_frame, text="Receipt preview")
        self.receipt_label.grid(column=0, row=0)

        self.receipt_frame.pack(side=tk.RIGHT)
        self.receipt_display = ttk.Treeview(
            self.receipt_frame,
            columns=(1,2,3),
            show="headings",
            height=10
        )
        self.receipt_display.grid(column=0, row=1)
        self.receipt_display.heading(1, text="Amount")
        self.receipt_display.column(1, anchor=tk.E)
        self.receipt_display.heading(2, text="Item")
        self.receipt_display.column(2, anchor=tk.W)
        self.receipt_display.heading(3, text="Price")
        self.receipt_display.column(3, anchor=tk.E)
        self.receipt_display.bind('<Double-1>', self.remove_item)
        
        # Receipt name
        self.receipt_name_frame = tk.Frame(self.receipt_frame)
        self.receipt_name_frame.grid(column=0, row=2)
        self.name_label = tk.Label(self.receipt_name_frame, text="Customer name")
        self.name_label.pack(side=tk.LEFT)
        self.name_entry = tk.Entry(self.receipt_name_frame, textvariable=self.customer_name)
        self.name_entry.pack(side=tk.RIGHT)
        
    # Items
    def load_items_from_workbook(self):
        wb_path = filedialog.askopenfilename(filetypes=[("Excel files",".xlsx")])
        print(f"Loading items from \"{wb_path}\"...")
        wb = load_workbook(wb_path)
        ws = wb.active
        for item, price in ws.iter_rows(min_row=2, min_col=1, max_col=2):
            self.item_dict.update({item.value: price.value})
        print(self.item_dict)
        for i, (item, price) in enumerate(self.item_dict.items()):
            self.item_list_display.insert(parent='', index=i, iid=i, values=(item, price))

    def add_item(self, event):
        # TODO: Support multiple items at the same time
        region = self.receipt_display.identify("region", event.x, event.y)
        if region == "heading":
            return

        item_index = int(self.item_list_display.selection()[0])
        item_list = list(self.item_dict.items())[item_index]
        item, price = item_list[0], item_list[1]
        print(f"Adding item #{item_index}: {item}")
        if item in self.receipt_items:
            self.receipt_items[item] += 1
        else:
            self.receipt_items.update({item: 1})

        self.refresh_receipt_preview()

    def remove_item(self, event):
        region = self.receipt_display.identify("region", event.x, event.y)
        if region == "heading":
            return

        item_index = int(self.receipt_display.selection()[0])
        item_list = list(self.receipt_items.items())[item_index]
        item = item_list[0]
        print(f"Removing item #{item_index}: {item}")
        self.receipt_items[item] -= 1
        if self.receipt_items[item] == 0:
            self.receipt_items.pop(item)
        self.refresh_receipt_preview()
        if len(self.receipt_items) > item_index:
            self.receipt_display.selection_add(item_index)

    def refresh_receipt_preview(self):
        for item in self.receipt_display.get_children():
            self.receipt_display.delete(item)
        
        self.receipt_items_for_docx = []
        prices = []
        for i, item in enumerate(self.receipt_items.items()):
            item_name = item[0]
            item_amount = item[1]
            item_price = item[1]*self.item_dict[item[0]]
            self.receipt_display.insert(
                parent='',index=i, iid=i,
                values=(item_amount, item_name, item_price)
                )
            self.receipt_items_for_docx.append((item_amount, item_name, item_price))
            prices.append(item_price)
        self.total_price = sum(prices)
        self.receipt_display.insert(
            parent='',index=i+1, iid=i+1,
            values=("", "", "----------")
            )
        self.receipt_items_for_docx.append(("", "", "----------"))
        self.receipt_display.insert(
            parent='',index=i+2, iid=i+2,
            values=("", "", f"P {self.total_price}")
            )
        self.receipt_items_for_docx.append(("", "", f"P {self.total_price}"))
        print(self.receipt_items)
        pass

    # def update_customer_name(self):
    #     self.customer_name.set()

    # Export to document
    def export_to_document(self):
        print("Exporting document...")
        cust_name = self.customer_name.get()
        document = Document()
        cust_name_p = document.add_paragraph(cust_name)
        cust_name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table = document.add_table(rows=len(self.receipt_items)+2, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        for row, data in enumerate(self.receipt_items_for_docx):
            for column, subdata in enumerate(data):
                table.cell(row, column).text = str(subdata)

        document.save(f"{datetime.now()} {cust_name}.docx")
    
    # Help
    def help(self):
        messagebox.showinfo(title="Help", message=self.HELP_MSG)

def main():
    root = tk.Tk()
    app = App(root)
    root.wm_title("Receipt maker")
    root.minsize(400, 200)
    root.resizable(False, False)
    app.mainloop()

if __name__ == "__main__":
    main()
